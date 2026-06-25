from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


class CoverPage:
    def __init__(
        self,
        judul: str,
        penulis: str,
        institusi: str,
        fakultas: str,
        nim: str = "",
        prodi: str = "",
        angkatan: str = "",
        tanggal: str = "",
        pembimbing: str = "",
    ):
        self.judul = judul
        self.penulis = penulis
        self.nim = nim
        self.institusi = institusi
        self.fakultas = fakultas
        self.prodi = prodi
        self.angkatan = angkatan
        self.tanggal = tanggal
        self.pembimbing = pembimbing

    def render(self, doc):
        cfg = doc.config
        document = doc.document

        # Remove first default paragraph
        if document.paragraphs:
            p = document.paragraphs[0]
            p.clear()

        # Add vertical spacing
        for _ in range(4):
            sp = document.add_paragraph()
            sp.paragraph_format.space_before = Pt(0)
            sp.paragraph_format.space_after = Pt(0)

        # Institution name
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(self.institusi.upper())
        run.font.name = cfg.font_name
        run.font.size = Pt(16)
        run.bold = True

        # Faculty
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(self.fakultas)
        run.font.name = cfg.font_name
        run.font.size = Pt(14)
        run.bold = True

        # Study program
        if self.prodi:
            p = document.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(self.prodi)
            run.font.name = cfg.font_name
            run.font.size = Pt(13)

        # Spacing
        document.add_paragraph()

        # Title
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(12)
        run = p.add_run(self.judul.upper())
        run.font.name = cfg.font_name
        run.font.size = Pt(16)
        run.bold = True

        # Spacing
        for _ in range(3):
            document.add_paragraph()

        # Author info
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"Nama Penulis: {self.penulis}")
        run.font.name = cfg.font_name
        run.font.size = Pt(12)

        if self.nim:
            p = document.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(f"NIM: {self.nim}")
            run.font.name = cfg.font_name
            run.font.size = Pt(12)

        if self.angkatan:
            p = document.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(f"Angkatan: {self.angkatan}")
            run.font.name = cfg.font_name
            run.font.size = Pt(12)

        if self.pembimbing:
            p = document.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(f"Pembimbing: {self.pembimbing}")
            run.font.name = cfg.font_name
            run.font.size = Pt(12)

        # Spacing
        for _ in range(2):
            document.add_paragraph()

        # Year
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(self.tanggal or "2026")
        run.font.name = cfg.font_name
        run.font.size = Pt(12)

        doc.add_page_break()
