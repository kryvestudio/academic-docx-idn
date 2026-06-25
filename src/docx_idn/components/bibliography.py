from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx_idn.core.document import _set_run_font


class Bibliography:
    def __init__(self, entries: list[dict] | None = None):
        self.entries = entries or []

    def render(self, doc):
        # Add heading with proper font
        heading = doc.document.add_heading(level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        heading.paragraph_format.space_before = Pt(24)
        heading.paragraph_format.space_after = Pt(12)
        heading.paragraph_format.line_spacing = 1.5
        heading.paragraph_format.first_line_indent = Cm(0)

        run = heading.add_run("DAFTAR PUSTAKA")
        _set_run_font(run, size=14, bold=True)

        for entry in self.entries:
            text = self._format_entry(entry)
            p = doc.document.add_paragraph()
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.left_indent = Cm(1.25)
            p.paragraph_format.first_line_indent = Cm(-1.25)
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            run = p.add_run(text)
            _set_run_font(run, size=12, bold=False)

    def _format_entry(self, entry: dict) -> str:
        parts = []
        penulis = entry.get("penulis", "")
        if penulis:
            parts.append(penulis)

        tahun = entry.get("tahun", "")
        if tahun:
            parts.append(f"({tahun})")

        judul = entry.get("judul", "")
        if judul:
            parts.append(f"{judul}.")

        edisi = entry.get("edisi", "")
        if edisi:
            parts.append(f"Edisi {edisi}.")

        penerbit = entry.get("penerbit", "")
        if penerbit:
            parts.append(f"{penerbit}.")

        url = entry.get("url", "")
        if url:
            parts.append(f"Diperoleh dari {url}")

        return " ".join(parts)
