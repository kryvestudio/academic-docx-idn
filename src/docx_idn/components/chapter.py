from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx_idn.core.document import _set_run_font
import re

ROMAN_NUMERALS = {
    1: "I", 2: "II", 3: "III", 4: "IV", 5: "V",
    6: "VI", 7: "VII", 8: "VIII", 9: "IX", 10: "X",
}


class Chapter:
    def __init__(self, judul: str, sections: list[dict] | None = None):
        self.judul = judul
        self.sections = sections or []

    def render(self, doc, number: int):
        roman = ROMAN_NUMERALS.get(number, str(number))

        # Create heading with two lines: "BAB II" and "TINJAUAN PUSTAKA"
        heading = doc.document.add_heading(level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        heading.paragraph_format.space_before = Pt(24)
        heading.paragraph_format.space_after = Pt(12)
        heading.paragraph_format.line_spacing = 1.5
        heading.paragraph_format.first_line_indent = Cm(0)

        # Line 1: "BAB II"
        run1 = heading.add_run(f"BAB {roman}")
        _set_run_font(run1, size=14, bold=True)

        # Line break
        run_br = heading.add_run()
        _set_run_font(run_br, size=14, bold=True)
        br = OxmlElement("w:br")
        run_br._r.append(br)

        # Line 2: "TINJAUAN PUSTAKA"
        run2 = heading.add_run(self.judul.upper())
        _set_run_font(run2, size=14, bold=True)

        # Render sections with numbering
        for i, section in enumerate(self.sections, start=1):
            self._render_section(doc, section, f"{number}.{i}")

    def _render_section(self, doc, section: dict, number: str):
        judul = section.get("judul")
        if not judul:
            raise ValueError(f"Section '{number}' tidak punya key 'judul'")
        label = f"{number} {judul}"

        # Create heading with proper font
        heading = doc.document.add_heading(level=2)
        heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        heading.paragraph_format.space_before = Pt(12)
        heading.paragraph_format.space_after = Pt(6)
        heading.paragraph_format.line_spacing = 1.5
        heading.paragraph_format.first_line_indent = Cm(0)

        run = heading.add_run(label)
        _set_run_font(run, size=12, bold=True)

        # Render body paragraphs (level 2 = sub-bab)
        for text in section.get("isi", []):
            self._render_body_paragraph(doc, text, level=2)

        # Render subsections (level 3) with numbering
        for j, sub in enumerate(section.get("subsections", []), start=1):
            self._render_subsection(doc, sub, f"{number}.{j}")

    def _render_subsection(self, doc, subsection: dict, number: str):
        judul = subsection.get("judul")
        if not judul:
            raise ValueError(f"Subsection '{number}' tidak punya key 'judul'")
        label = f"{number} {judul}"

        # Create heading with proper font
        heading = doc.document.add_heading(level=3)
        heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        heading.paragraph_format.left_indent = Cm(1.25)
        heading.paragraph_format.space_before = Pt(6)
        heading.paragraph_format.space_after = Pt(6)
        heading.paragraph_format.line_spacing = 1.5
        heading.paragraph_format.first_line_indent = Cm(0)

        run = heading.add_run(label)
        _set_run_font(run, size=12, bold=True)

        # Render body paragraphs (level 3 = sub-sub-bab)
        for text in subsection.get("isi", []):
            self._render_body_paragraph(doc, text, level=3)

    def _render_body_paragraph(self, doc, text: str, level: int = 2):
        """Render body paragraph with proper formatting and hanging indent for numbered lists."""
        p = doc.document.add_paragraph()
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        # Check if text starts with a number (like "1. ", "2. ")
        match = re.match(r'^(\d+\.\s*)', text)
        if match:
            # Hanging indent for numbered lists
            p.paragraph_format.left_indent = Cm(1.25)
            p.paragraph_format.first_line_indent = Cm(-1.25)
        elif level == 3:
            # Sub-sub-bab: deeper indent than sub-bab
            p.paragraph_format.left_indent = Cm(2.0)
            p.paragraph_format.first_line_indent = Cm(0)
        else:
            # Normal paragraph (sub-bab level)
            p.paragraph_format.first_line_indent = Cm(1.25)

        run = p.add_run(text)
        _set_run_font(run, size=12, bold=False)
