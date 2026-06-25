from docx import Document
from docx.shared import Pt
from docx_idn.core.config import DocumentConfig
from docx_idn.core.style import StyleManager
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def _set_run_font(run, font_name="Times New Roman", size=12, bold=False):
    """Set font for a run including East Asian font."""
    run.font.name = font_name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = None
    # Set East Asian font
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)
    rFonts.set(qn("w:eastAsia"), font_name)
    rFonts.set(qn("w:cs"), font_name)


class DocxDocument:
    def __init__(self, config: DocumentConfig | None = None):
        self.config = config or DocumentConfig()
        self.document = Document()
        self.style_manager = StyleManager(self.document, self.config)
        self.style_manager.setup_page()
        self.style_manager.setup_styles()

    def add_paragraph(self, text: str, style: str | None = None):
        p = self.document.add_paragraph(text, style=style)
        if style is None:
            self.style_manager.setup_default_paragraph(p)
        return p

    def add_heading(self, text: str, level: int = 1):
        """Add heading with proper Times New Roman font."""
        heading = self.document.add_heading(level=level)
        # Set font size based on level
        size_map = {1: self.config.font_size_h1, 2: self.config.font_size_h2, 3: self.config.font_size_h3}
        size = size_map.get(level, self.config.font_size)
        run = heading.add_run(text)
        _set_run_font(run, size=size, bold=True)
        return heading

    def add_page_break(self):
        self.document.add_page_break()

    def save(self, path: str):
        self.document.save(path)
