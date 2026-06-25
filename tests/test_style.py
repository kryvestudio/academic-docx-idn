from docx import Document
from docx_idn.core.style import StyleManager
from docx_idn.core.config import DocumentConfig
from docx.shared import Cm


def test_setup_page():
    doc = Document()
    cfg = DocumentConfig()
    sm = StyleManager(doc, cfg)
    sm.setup_page()
    section = doc.sections[0]
    # Allow small rounding tolerance (within 1mm)
    assert abs(section.left_margin - Cm(4.0)) < Cm(0.1)
    assert abs(section.right_margin - Cm(3.0)) < Cm(0.1)
    assert abs(section.top_margin - Cm(4.0)) < Cm(0.1)
    assert abs(section.bottom_margin - Cm(3.0)) < Cm(0.1)


def test_normal_style():
    doc = Document()
    cfg = DocumentConfig()
    sm = StyleManager(doc, cfg)
    sm.setup_styles()
    style = doc.styles["Normal"]
    font = style.font
    assert font.name == "Times New Roman"
    assert font.size.pt == 12
    pf = style.paragraph_format
    assert pf.line_spacing == 1.5


def test_heading1_style():
    doc = Document()
    cfg = DocumentConfig()
    sm = StyleManager(doc, cfg)
    sm.setup_styles()
    style = doc.styles["Heading 1"]
    font = style.font
    assert font.name == "Times New Roman"
    assert font.size.pt == 14
    assert font.bold is True
