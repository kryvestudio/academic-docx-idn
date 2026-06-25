from docx import Document
from docx_idn.core.document import DocxDocument
from docx_idn.core.config import DocumentConfig


def test_docx_document_creates_doc():
    doc = DocxDocument(DocumentConfig())
    assert doc.document is not None


def test_docx_document_saves(tmp_path):
    doc = DocxDocument(DocumentConfig())
    doc.add_paragraph("Test")
    out = tmp_path / "test.docx"
    doc.save(str(out))
    assert out.exists()
